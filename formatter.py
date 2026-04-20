"""Formatter agent for STOFabric stage 5.

Builds final DOCX from transformed STO JSON.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


class Formatter:
    """Formats transformed STO JSON into a final DOCX document."""

    def __init__(self, transformed_payload: dict[str, Any]) -> None:
        self.payload = transformed_payload
        self.sto_document = transformed_payload.get("sto_document_json", transformed_payload)
        self.warnings: list[str] = []
        self.doc = Document()
        self.table_counter = 0
        self.figure_counter = 0
        self._configure_document_defaults()

    @staticmethod
    def load_json(path: str | Path) -> dict[str, Any]:
        return json.loads(Path(path).read_text(encoding="utf-8"))

    def _configure_document_defaults(self) -> None:
        """Apply baseline page and text settings."""
        section = self.doc.sections[0]
        section.left_margin = Mm(20)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.right_margin = Mm(10)
        section.different_first_page_header_footer = True

        style = self.doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.first_line_indent = Mm(12.5)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _add_page_number_field(self, paragraph) -> None:
        """Insert PAGE field to paragraph."""
        run = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)

    def _set_header_footer(self) -> None:
        sto_number = self.sto_document.get("meta", {}).get("sto_number", "")
        for section in self.doc.sections:
            section.different_first_page_header_footer = True
            header_par = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
            header_par.text = sto_number
            header_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

            footer_par = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
            footer_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_page_number_field(footer_par)

    def build_title_page(self) -> None:
        meta = self.sto_document.get("meta", {})
        approval = meta.get("approval", {})
        org_name = meta.get("organization", {}).get("name", "ООО «ПКФ «СНАРК»")
        self.doc.add_paragraph(org_name).alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph("СТАНДАРТ ОРГАНИЗАЦИИ").alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph(meta.get("sto_number", "СТО 31025229-000-2024")).alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p = self.doc.add_paragraph(meta.get("title", "Стандарт организации"), style="Title")
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.runs[0].bold = True
        self.doc.add_paragraph()
        self.doc.add_paragraph("УТВЕРЖДАЮ")
        self.doc.add_paragraph(approval.get("approver_position", "Генеральный директор"))
        self.doc.add_paragraph(approval.get("approver_name", ""))
        self.doc.add_paragraph(str(approval.get("approval_date", meta.get("approval_date", ""))))
        self.doc.add_page_break()

    def _collect_main_toc_entries(self, sections: list[dict[str, Any]]) -> list[str]:
        entries: list[str] = []
        for section in sections:
            number = section.get("section_number", "").strip()
            title = section.get("title", "").strip()
            line = f"{number} {title}".strip()
            if line:
                entries.append(line)
            entries.extend(self._collect_main_toc_entries(section.get("subsections", [])))
        return entries

    def build_toc(self) -> None:
        self.doc.add_paragraph("Содержание", style="Heading 1").runs[0].bold = True
        toc_items = [
            "1 Область применения",
            "2 Нормативные ссылки",
            "3 Термины, определения и сокращения",
            "4 Основная часть",
            "5 Отчетные документы",
            "6 Ответственность",
        ]
        toc_items.extend(self._collect_main_toc_entries(self.sto_document.get("content", {}).get("main", [])))
        if self.sto_document.get("appendices"):
            toc_items.append("Приложения")
        toc_items.extend(["Лист регистрации изменений", "Лист ознакомления", "Лист согласования"])
        for item in toc_items:
            self.doc.add_paragraph(item)
        self.doc.add_page_break()

    def _add_section_heading(self, title: str, level: int = 1) -> None:
        style_name = "Heading 1"
        if level == 2:
            style_name = "Heading 2"
        elif level >= 3:
            style_name = "Heading 3"
        p = self.doc.add_paragraph(title, style=style_name)
        if p.runs:
            p.runs[0].bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _write_tables(self, tables: list[dict[str, Any]]) -> None:
        for table in tables:
            self.table_counter += 1
            title = table.get("title", f"Таблица {self.table_counter}")
            self.doc.add_paragraph(f"Таблица {self.table_counter} — {title}")
            rows = table.get("rows", [])
            if not rows:
                continue
            col_count = max((len(r.get("cells", [])) for r in rows), default=1)
            doc_table = self.doc.add_table(rows=len(rows), cols=col_count)
            doc_table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                cells = row.get("cells", [])
                for c_idx in range(col_count):
                    doc_table.cell(r_idx, c_idx).text = str(cells[c_idx]) if c_idx < len(cells) else ""
            self.doc.add_paragraph()

    def _extract_image_path(self, image: dict[str, Any]) -> str | None:
        if image.get("path"):
            return str(image["path"])
        image_ref = image.get("image_ref")
        if isinstance(image_ref, list):
            for item in image_ref:
                if isinstance(item, dict) and item.get("path"):
                    return str(item["path"])
        return None

    def _write_images(self, images: list[dict[str, Any]]) -> None:
        for image in images:
            self.figure_counter += 1
            caption = image.get("caption") or image.get("title") or f"Рисунок {self.figure_counter}"
            self.doc.add_paragraph(f"Рисунок {self.figure_counter} — {caption}")
            path = self._extract_image_path(image)
            if path:
                try:
                    self.doc.add_picture(path)
                    continue
                except Exception:
                    self.warnings.append(f"Image file is unavailable: {path}")
            self.doc.add_paragraph("[Изображение отсутствует в источнике]")

    def build_mandatory_sections(self) -> None:
        content = self.sto_document.get("content", {})
        self._add_section_heading("1 Область применения", level=1)
        self.doc.add_paragraph(content.get("area", ""))

        self._add_section_heading("2 Нормативные ссылки", level=1)
        refs = content.get("normative_references", [])
        if refs:
            for ref in refs:
                self.doc.add_paragraph(f"{ref.get('reference_id', '')} — {ref.get('title', '')}")

        self._add_section_heading("3 Термины, определения и сокращения", level=1)
        terms = content.get("terms", {})
        if terms.get("intro"):
            self.doc.add_paragraph(terms["intro"])
        for item in terms.get("entries", []):
            self.doc.add_paragraph(f"{item.get('term', '')}: {item.get('definition', '')}")
        for ab in terms.get("abbreviations", []):
            self.doc.add_paragraph(f"{ab.get('abbr', '')} - {ab.get('definition', '')}")

    def _build_main_sections_recursive(self, sections: list[dict[str, Any]], level: int = 2) -> None:
        for section in sections:
            title = f"{section.get('section_number', '')} {section.get('title', '')}".strip()
            self._add_section_heading(title, level=level)
            if section.get("content"):
                self.doc.add_paragraph(section["content"])
            for point in section.get("bullet_points", []):
                self.doc.add_paragraph(f"• {point}")
            self._write_tables(section.get("tables", []))
            self._write_images(section.get("images", []))
            self._build_main_sections_recursive(section.get("subsections", []), level=min(level + 1, 3))

    def build_main_sections(self) -> None:
        self._add_section_heading("4 Основная часть", level=1)
        self._build_main_sections_recursive(self.sto_document.get("content", {}).get("main", []), level=2)

    def build_reporting_responsibility_appendices(self) -> None:
        self._add_section_heading("5 Отчетные документы", level=1)
        for row in self.sto_document.get("reporting_documents", []):
            self.doc.add_paragraph(
                " | ".join(
                    [
                        row.get("document_name", ""),
                        row.get("responsible_role", ""),
                        row.get("storage_location", ""),
                        row.get("retention_period", ""),
                    ]
                )
            )

        self._add_section_heading("6 Ответственность", level=1)
        resp = self.sto_document.get("responsibility", {}).get("entries", [])
        for entry in resp:
            self.doc.add_paragraph(entry.get("role", ""))
            for item in entry.get("responsibilities", []):
                self.doc.add_paragraph(f"- {item}")

        for appendix in self.sto_document.get("appendices", []):
            self.doc.add_page_break()
            app_header = f"Приложение {appendix.get('appendix_id', '')} ({appendix.get('status', 'справочное')})"
            self._add_section_heading(app_header, level=1)
            self.doc.add_paragraph(appendix.get("title", ""))
            if appendix.get("content_text"):
                self.doc.add_paragraph(appendix["content_text"])
            self._write_tables(appendix.get("tables", []))
            self._write_images(appendix.get("images", []))

    def build_service_sheets(self) -> None:
        sheets = self.sto_document.get("meta", {}).get("extra_attributes", {}).get("service_sheets", {})
        self.doc.add_page_break()
        self._add_section_heading("Лист регистрации изменений", level=1)
        self._render_service_sheet_table(
            sheets.get("change_log", []),
            ["Номер изменения", "Номер извещения", "Дата", "Подпись"],
        )

        self.doc.add_page_break()
        self._add_section_heading("Лист ознакомления", level=1)
        self._render_service_sheet_table(
            sheets.get("acquaintance", []),
            ["№", "ФИО", "Должность", "Подпись", "Дата"],
        )

        self.doc.add_page_break()
        self._add_section_heading("Лист согласования", level=1)
        self._render_service_sheet_table(
            sheets.get("approval", []),
            ["Должность", "ФИО", "Подпись", "Дата"],
        )

    def _render_service_sheet_table(self, items: list[dict[str, Any]], columns: list[str]) -> None:
        table = self.doc.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"
        for idx, col in enumerate(columns):
            table.cell(0, idx).text = col
            table.cell(0, idx).paragraphs[0].runs[0].bold = True
        rows_to_render = items if items else [{} for _ in range(5)]
        for item in rows_to_render:
            row = table.add_row().cells
            values = [str(item.get(col, item.get(col.lower(), ""))) for col in columns]
            for idx, value in enumerate(values):
                row[idx].text = value

    def apply_global_styles_and_layout(self) -> None:
        self._set_header_footer()

    def save_docx(self, output_path: str | Path) -> None:
        self.doc.save(str(output_path))

    def build(self, output_path: str | Path) -> dict[str, Any]:
        self.build_title_page()
        self.build_toc()
        self.build_mandatory_sections()
        self.build_main_sections()
        self.build_reporting_responsibility_appendices()
        self.build_service_sheets()
        self.apply_global_styles_and_layout()
        self.save_docx(output_path)
        return {"output_docx": str(output_path), "warnings": self.warnings}


def _build_cli() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Format transformed STO JSON to DOCX")
    parser.add_argument("--input-json", required=True, type=Path, help="Transformed JSON path")
    parser.add_argument("--output-docx", required=True, type=Path, help="Output DOCX path")
    return parser


def main() -> None:
    args = _build_cli().parse_args()
    try:
        payload = Formatter.load_json(args.input_json)
        formatter = Formatter(payload)
        result = formatter.build(args.output_docx)
        print(json.dumps(result, ensure_ascii=True, indent=2))
    except Exception as exc:  # pragma: no cover
        print(json.dumps({"error": str(exc)}, ensure_ascii=False, indent=2))
        raise SystemExit(1) from exc


if __name__ == "__main__":
    main()
